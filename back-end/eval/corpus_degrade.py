"""语料降级：把干净的 Markdown 变成"真实上传"的样子。

这个模块存在的理由是一个评估集盲区：``runner.ensure_corpus`` 只收 ``.md``，而
``eval/corpus/`` 下那 6 篇是自造的、干净的、utf-8 的 Markdown。**于是这套评估
量不出任何清洗改动的价值**——清洗代码在这条链路上根本没有输入可清。

补一份真实 PDF 夹具解决不了这个问题，反而更糟：它是一个不透明的数据点，脏在
哪一处、哪一处损伤伤到了召回，全都分不开；而且要往仓库里提交二进制。这里改为
**具名、确定性、可组合的降级**，每一种只模拟一类真实故障，于是"清洗值多少"能
拆到每一类损伤上。

三条设计约束：

1. **确定性。** 不用随机数，全按行号/字符位取模。同一份语料降两次逐位相同，
   否则同一个变体跑两遍得到不同的召回，变体差异被方差盖掉——这和 ``agent_runner``
   把温度钉在 0.0 是同一个理由。
2. **只动形式，不动事实。** 金标准里的 ``must_include``、``expected_documents``
   全部复用，所以降级不能删掉正文里的数字、术语、编号。``pdf_like`` 抹掉的是
   ``#`` 标记和空行，不是标题文字本身。
3. **降级名必须进分块指纹。** 见 ``runner._chunking_fingerprint``：不并入的话
   第二个变体会命中上一个变体留下的索引，测的是上一次的配置。

诚实的局限：``pdf_like`` 是合成损伤，不等于真实 PyPDF2 的输出——真实抽取还会
打乱多栏顺序、把表格拍平成一行。``gbk_bytes`` 与 ``noisy_unicode`` 是真的
（真的重新编码、真的插入那些码位）。

**两类降级测的不是同一个问题，别混着读：**

- ``pdf_like`` 的损伤**清洗修不了**，这是设计如此：词内空格、丢掉的 ``#`` 标记、
  页眉页脚，全都只能靠 PDF 的字号与坐标复原，而降级产物是纯文本，没有几何信息。
  所以它没有 ``+clean`` 对照组，它回答的是另一个问题——**丢掉结构要付多少代价**，
  也就是"第 1 条（PDF 结构恢复）值不值得做"。和 ``baseline`` 比。
- ``gbk_bytes`` / ``noisy_unicode`` 的损伤清洗**修得了**，所以它们成对存在，
  差值就是"清洗追回了多少"。
"""
from __future__ import annotations

import io
import re

# 页眉页脚每隔这么多行注入一次。真实 PDF 是每页一次，而纯文本没有页概念，
# 用固定行距模拟——重要的是"跨页重复出现"这个特征被保留下来，因为
# ingest_clean._strip_running_heads 的判据正是跨页重复。
_PAGE_LINES = 18
_RUNNING_HEAD = "公司内部资料 · 未经许可不得外传"
# 词内空格的插入步长。每 N 个拉丁/数字字符插一个空格，模拟 PDF 抽取按字距
# 切词切错的效果——这一条直接打在 BM25 的词元上。
_WORD_BREAK_STRIDE = 7
_LATIN_RUN_RE = re.compile(r"[A-Za-z0-9_]{5,}")
_HEADING_RE = re.compile(r"^(#{1,6})\s+(.*?)\s*#*$")

# noisy_unicode 用的：零宽字符按固定步长插进长词里
_ZERO_WIDTH_CHARS = ("​", "­", "‌")
_ZERO_WIDTH_STRIDE = 6
_FULLWIDTH_OFFSET = 0xFF01 - 0x21

DEGRADATIONS = (
    "none",
    "pdf_like",
    "gbk_bytes",
    "noisy_unicode",
    "scanned",
    # 严格说 docx 不是"降级"——它是格式转换，而且几乎无损。放在这个模块里是因为
    # 机制完全一样（同一份语料换一种上传形态、名字进指纹、金标不动），
    # 而为它另开一个"格式变体"维度会让 ensure_corpus 多一条正交的分支。
    "docx",
)

# 为什么**没有** xlsx 模式（2026-08-24 评估过，四条都是硬的）：
#
# 1. 13 篇语料只有 9 篇含表格 → 另外 4 篇会变成空文档。
# 2. 即使只看 table_lookup 那 6 条金标，其中 1 条（"8000 元报销要谁审批" → 财务）
#    的答案在**无序列表**里而不是表格里（expense-policy.md 的审批门槛那段）。
# 3. 剥掉散文只留表格之后，30 条金标里约 24 条**按设计必然失败**，
#    报告是一片零，只剩一个很弱的信号。
# 4. 诚实的设计是"散文留 .md、表格拆成 .xlsx"，但 degrade_corpus_file 的契约是
#    **1 篇源 → 1 篇上传**，而 ensure_corpus 的早退条件正是 len(文档)==len(源文件)。
#    拆分要改这个契约，那是另一件事。
#
# xlsx 解析本身已有 29 条单元测试覆盖（tests/test_ingest_xlsx.py），包括"每个分块
# 都自足"这条核心取舍以及它的 Markdown 表格对照组。eval 能**额外**给的只有一件事：
# "``列名: 值`` 这种行在真实检索下召不召得回来"。那确实有价值，但它需要表格形态的
# 语料内容，而造那批内容等于写新金标——也就是改考题，不是测能力。
#
# 真要做的前置条件：一批**本来就是表格**的语料（费率表、额度表）+ 配套金标，
# 且 degrade_corpus_file 支持 1 源 → N 上传。

# Markdown 表格的分隔行：``| --- | :--: |``。判据是去掉管道之后只剩 - : 和空白。
_TABLE_DIVIDER_RE = re.compile(r"^\s*\|[\s\-:|]+\|\s*$")
_TABLE_ROW_RE = re.compile(r"^\s*\|(.+)\|\s*$")
# 无序列表标记。docx 里渲染成 List Bullet 段落而不是保留 `- `：
# Word 文档里的列表就是这个形状，而 extract_docx 读回来是纯段落文本。
_LIST_ITEM_RE = re.compile(r"^\s*[-*+]\s+(.*)$")


def _strip_markdown_structure(text: str) -> str:
    """删掉 ``#`` 标记、列表标记之外的结构，保留文字。

    这一条是 PDF 损伤里最贵的：``.pdf`` 不在 ``chunking._MARKDOWN_EXTENSIONS``
    里，抽出的文本也没有 ``#`` 行，于是 ``_looks_like_markdown`` 判 False →
    ``heading_path`` 恒为空 → 「给每块加标题路径」和「章节边界优先于填满
    max_tokens」两件事全部失效。标题**文字**留着，所以 must_include 仍然可达。
    """
    lines = []
    for line in text.splitlines():
        match = _HEADING_RE.match(line)
        lines.append(match.group(2) if match else line)
    return "\n".join(lines)


def _break_words(text: str) -> str:
    """在长拉丁/数字词里按固定步长插空格。

    打的是 ``RESOURCE_EXHAUSTED`` → ``RESOURCE _EXHAUS TED`` 这类损伤，也就是
    金标准里 6 条 ``lexical`` 探针考的那种字面命中。中文不动：真实 PDF 抽取
    对中文的典型损伤是**丢**空格结构而不是插空格，而且往中文里插空格会让
    tokenize 的 CJK bigram 全废，那是另一类损伤（``none`` 之外没有变体测它）。
    """

    def split_run(match: re.Match[str]) -> str:
        run = match.group(0)
        pieces = [
            run[start : start + _WORD_BREAK_STRIDE]
            for start in range(0, len(run), _WORD_BREAK_STRIDE)
        ]
        return " ".join(pieces)

    return _LATIN_RUN_RE.sub(split_run, text)


def _inject_running_heads(text: str) -> str:
    """每隔固定行数插一次页眉和页码。

    页码故意用递增数字：``_strip_running_heads`` 靠把数字折成 ``#`` 才能把
    「第 1 页」和「第 12 页」collapse 成同一个频率键，不递增就测不到那一步。
    """
    lines = text.splitlines()
    out: list[str] = []
    page = 1
    for index, line in enumerate(lines):
        if index and index % _PAGE_LINES == 0:
            out.append(f"第 {page} 页")
            page += 1
            out.append(_RUNNING_HEAD)
        out.append(line)
    return "\n".join(out)


def _collapse_blank_lines(text: str) -> str:
    """删掉空行。

    ``chunking._parse_blocks`` 靠空行 flush 段落，没有空行整个章节就变成一个
    大 block，只能靠 ``_pack_units`` 按句硬切——段落边界这个信息彻底没了。
    """
    return "\n".join(line for line in text.splitlines() if line.strip())


def pdf_like(text: str) -> str:
    """模拟 PyPDF2 抽取 PDF 的输出。四种损伤叠加，顺序固定。

    这些损伤**清洗修不了**（纯文本里没有字号和坐标），所以它没有 ``+clean``
    对照组。它量的是"丢掉结构要付多少代价"。
    """
    damaged = _strip_markdown_structure(text)
    damaged = _inject_running_heads(damaged)
    damaged = _break_words(damaged)
    return _collapse_blank_lines(damaged)


def noisy_unicode(text: str) -> str:
    """全角 ASCII + 词内零宽字符 + CRLF + 控制字符。

    这三种都是真实文档里极常见的东西：Word / PDF 导出会把半角数字变成全角，
    从网页复制粘贴会带进零宽空格和软连字符，Windows 编辑器留 CRLF。

    它们比 ``pdf_like`` 隐蔽得多，因为**看起来完全正常**——``４２９`` 和 ``429``
    在屏幕上几乎一样，零宽字符根本看不见。但在 ``tokenize()`` 眼里
    ``４２９`` 不是 ``429``，夹了 U+200B 的 ``RESOURCE_EXHAUSTED`` 是两个词元。
    而这一类正好是 ``clean_text`` 修得了的，所以它有 ``+clean`` 对照组。
    """
    folded = []
    for char in text:
        code = ord(char)
        # 只全角化 ASCII 可见字符里的字母数字，标点不动：全角标点是中文文档里
        # 的正常写法，把它算成"损伤"会让这个降级测的东西变得不干净
        if char.isalnum() and 0x21 <= code <= 0x7E:
            folded.append(chr(code + _FULLWIDTH_OFFSET))
        else:
            folded.append(char)
    noisy = "".join(folded)

    def inject(match: re.Match[str]) -> str:
        run = match.group(0)
        pieces = [
            run[start : start + _ZERO_WIDTH_STRIDE]
            for start in range(0, len(run), _ZERO_WIDTH_STRIDE)
        ]
        # 取模而不是随机：同一个词每次都插同一个零宽字符
        return "".join(
            piece + (_ZERO_WIDTH_CHARS[index % len(_ZERO_WIDTH_CHARS)] if index < len(pieces) - 1 else "")
            for index, piece in enumerate(pieces)
        )

    # 全角化之后原来的拉丁串已经不匹配 _LATIN_RUN_RE 了，所以零宽注入按
    # CJK 与全角字符的连续段来做——这也更贴近真实：中文正文里同样会夹零宽字符
    noisy = re.sub(r"[０-ｚ一-鿿]{6,}", inject, noisy)
    # 行尾统一成 CRLF，再夹几个控制字符
    lines = noisy.split("\n")
    return "\r\n".join(
        line + ("\x0c" if index and index % 12 == 0 else "")
        for index, line in enumerate(lines)
    )


def _parse_table_row(line: str) -> list[str] | None:
    """``| a | b |`` → ``["a", "b"]``。不是表格行就返回 None。"""
    match = _TABLE_ROW_RE.match(line)
    if not match:
        return None
    return [cell.strip() for cell in match.group(1).split("|")]


def markdown_to_docx(text: str) -> bytes:
    """把 Markdown 语料转成一份真 .docx。

    ## 这个模式量的是什么

    ``md → docx → extract_docx`` 是一条**近乎恒等**的往返：``## 标题`` 变成
    Heading 2 样式、再被读回成 ``## 标题``；Markdown 表格变成 Word 表格、
    再被读回成 Markdown 表格。所以这个变体的预期结果是**贴着 baseline**。

    这让它成为块 3 那个解析器的回归测试，而且是用检索指标表达的：
    差值接近零说明结构保真；掉下来就说明 ``extract_docx`` 丢了真东西
    （标题层级没识别、表格跑到文末、单元格换行没压平……），而那些症状在单元
    测试里都是"能读出文字"所以全绿。

    和 ``pdf_like`` 的关系正好相反：那个刻意丢掉结构，量"丢了值多少"；
    这个刻意保留结构，量"我们的解析器有没有真的把它保下来"。

    ## 为什么不用 pandoc

    要一个外部二进制，而降级必须是确定性的、在 CI 里可复现的。这里的转换只需要
    覆盖语料实际用到的四种构造（标题、段落、无序列表、表格），
    ``python-docx`` 直接写就够，也不引入版本漂移。

    ## 已知的不保真处

    行内标记（``**粗体**``、`` `代码` ``）原样进段落文本。语料里它们很少，
    而且 ``extract_docx`` 读回来也是同样的字面量，所以往返仍然一致——
    只是它没有变成 Word 的真粗体。这不影响 BM25 词元，故不修。
    """
    import docx  # 局部导入：只有这一个降级模式需要它

    document = docx.Document()
    pending_rows: list[list[str]] = []

    def flush_table() -> None:
        """把攒下的表格行写成一张 Word 表格。"""
        if not pending_rows:
            return
        width = max(len(row) for row in pending_rows)
        table = document.add_table(rows=len(pending_rows), cols=width)
        for row_index, cells in enumerate(pending_rows):
            padded = cells + [""] * (width - len(cells))
            for column_index, value in enumerate(padded):
                table.cell(row_index, column_index).text = value
        pending_rows.clear()

    for line in text.splitlines():
        stripped = line.strip()

        # 表格：连续的表格行攒起来一次性写，分隔行丢掉（Word 表格没有这个概念）
        if _TABLE_DIVIDER_RE.match(line):
            continue
        row = _parse_table_row(line) if stripped.startswith("|") else None
        if row is not None:
            pending_rows.append(row)
            continue
        flush_table()

        if not stripped:
            continue

        heading = _HEADING_RE.match(stripped)
        if heading:
            level = min(6, len(heading.group(1)))
            document.add_heading(heading.group(2), level=level)
            continue

        list_item = _LIST_ITEM_RE.match(line)
        if list_item:
            document.add_paragraph(list_item.group(1), style="List Bullet")
            continue

        document.add_paragraph(stripped)

    flush_table()

    buffer = io.BytesIO()
    document.save(buffer)
    return _normalize_zip_timestamps(buffer.getvalue())


# ZIP 条目时间戳统一成这个值。取一个固定的过去时间而不是 0：
# ZIP 的时间字段存不下 1980 年之前的日期，zipfile 会对 0 报警。
_FIXED_ZIP_DATETIME = (1980, 1, 1, 0, 0, 0)


def _normalize_zip_timestamps(payload: bytes) -> bytes:
    """把 .docx（一个 ZIP）里所有条目的时间戳抹平，让输出逐位可复现。

    没有这一步 ``markdown_to_docx`` 就**不满足这个模块的第一条设计约束**
    （"同一份语料降两次逐位相同"）：ZIP 条目带修改时间，精度 2 秒，所以
    跨过一个 2 秒边界的两次调用字节不同。

    这一条是实测出来的，不是预防性的：现成的
    ``test_every_degradation_is_deterministic`` 对 docx **本来是条 flake**——
    它只在两次调用落在同一个 2 秒窗口里才绿，而那是大多数时候。

    时间戳其实不影响检索（``_corpus_digest`` 摘的是源 ``.md``，而解析出来的
    正文与时间戳无关），所以也可以选择把那条测试改成"对 docx 只比正文"。
    没那么做：字节确定性本身是有用的（可复现产物、可缓存），而放宽一条已经写下
    的不变量，代价是以后没人知道它还成不成立。
    """
    import zipfile

    source = zipfile.ZipFile(io.BytesIO(payload))
    out = io.BytesIO()
    # 条目顺序沿用原样：ZIP 的中央目录顺序也进字节流，重排同样会破坏逐位相同
    with zipfile.ZipFile(out, "w", zipfile.ZIP_DEFLATED) as target:
        for item in source.infolist():
            normalized = zipfile.ZipInfo(item.filename, _FIXED_ZIP_DATETIME)
            normalized.compress_type = item.compress_type
            normalized.external_attr = item.external_attr
            normalized.internal_attr = item.internal_attr
            normalized.create_system = item.create_system
            target.writestr(normalized, source.read(item.filename))
    source.close()
    return out.getvalue()


def degrade_corpus_file(text: str, mode: str) -> tuple[bytes, str]:
    """按降级方式返回 (上传用的字节, 建议的文件后缀)。

    返回 bytes 而不是 str：``gbk_bytes`` 的全部意义就在于它**不是** utf-8，
    在字符串层面无法表达。后缀跟着变是因为 ``chunking._looks_like_markdown``
    会看扩展名——``pdf_like`` 的文本必须以非 ``.md`` 的名字进去，否则
    "PDF 没有标题层级"这个损伤会被扩展名兜回来，测出来的差值是假的。
    """
    if mode == "none":
        return text.encode("utf-8"), ".md"
    if mode == "pdf_like":
        return pdf_like(text).encode("utf-8"), ".txt"
    if mode == "gbk_bytes":
        # errors="replace" 让语料里偶发的生僻字不至于让整个降级失败;
        # 被替换的是极少数字符,不影响"整篇按 utf-8 解会烂掉"这个待测特征
        return text.encode("gbk", errors="replace"), ".md"
    if mode == "noisy_unicode":
        return noisy_unicode(text).encode("utf-8"), ".md"
    if mode == "scanned":
        # 图片型 PDF：有文件、有页数、抽不出任何文本
        return b"", ".txt"
    if mode == "docx":
        return markdown_to_docx(text), ".docx"
    raise ValueError(f"未知的降级方式: {mode}（可用: {', '.join(DEGRADATIONS)}）")
