"""摄取层清洗：解码、PDF 结构恢复、文本规整。

这个模块存在的理由不是"预处理让质量好一点"，而是一件具体得多的事：**脏输入会让
整条召回通道彻底失效，而且失效得完全没有声音**。三条都能在代码里对上：

1. GBK 文档按 utf-8 解码成一串 U+FFFD 之后，``retrieval_index.tokenize()`` 的两个
   正则（拉丁词元、CJK bigram）一个都匹配不到，BM25 建索引时 ``if not tokens:
   continue`` 直接跳过整块——**稀疏通道看不见这篇文档**。稠密通道拿到的也只是
   替换符的向量。而文档状态是 ``indexed``、``chunks=N``，界面上完全正常。
2. PDF 抽出来的文本没有 ``#`` 行，于是 ``chunking._looks_like_markdown`` 判 False，
   ``heading_path`` 恒为空：「给每块加标题路径」和「章节边界优先于填满
   ``max_tokens``」这两件事对 PDF 全都不生效。``chunking`` 模块文档里承诺的四件事
   只剩两件。
3. 抽取把一个词切成 ``RESOURCE_ EXHAUSTED``，bigram 与词元全变——而 BM25 通道
   存在的全部理由就是接住这类字面命中。

所以这里做的是把两条通道的输入契约补齐。

**为什么用 pdfplumber 而不是继续用 PyPDF2**：第 2 条只能靠字号和坐标修。
``PdfReader.extract_text()`` 返回的是一个字符串，里面既没有字号也没有位置，
拿它无论怎么正则都推不出标题层级。pdfplumber 的 ``extract_words`` 每个词带
``size`` / ``x0`` / ``top``，标题识别、页眉页脚定位、词内空格判定全依赖它们。
"""
from __future__ import annotations

import io
import logging
import re
from collections import Counter
from dataclasses import dataclass, field
from datetime import date, datetime
from typing import Any

from config import settings

logger = logging.getLogger("ingest_clean")

# ========== 文本规整 ==========

# 零宽字符与软连字符。它们在 tokenize() 眼里既不是拉丁词元也不是 CJK，
# 于是一个夹在词中间的 U+200B 就能让 `RESOURCE_EXHAUSTED` 变成两个词元。
_INVISIBLE = dict.fromkeys(
    [
        0x200B,  # 零宽空格
        0x200C,  # 零宽不连字
        0x200D,  # 零宽连字
        0x2060,  # word joiner
        0xFEFF,  # BOM 出现在正文中间时
        0x00AD,  # 软连字符
    ],
    None,
)
# 控制字符（保留 \t \n）。PDF 与某些导出工具会夹带 \x0c 分页符之类的东西
_CONTROL_RE = re.compile(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]")
_BLANK_RUN_RE = re.compile(r"\n{3,}")
_TRAILING_WS_RE = re.compile(r"[ \t]+$", re.MULTILINE)
_REPLACEMENT = "�"

# 全角 ASCII → 半角。只折这一段，不做整体 NFKC：NFKC 会顺手动到 CJK 兼容
# 字符（``㈱`` 之类）和一些标点，范围比这里需要的大得多。
# 折它的理由很实际：全角的 ４２９ 和半角的 429 在 tokenize() 里是**不同的词元**，
# 而金标准里 lexical 探针考的就是 429、P8 这类字面命中。
_FULLWIDTH_START = 0xFF01
_FULLWIDTH_END = 0xFF5E
_FULLWIDTH_OFFSET = 0xFF01 - 0x21
_IDEOGRAPHIC_SPACE = "　"


def _fold_fullwidth(text: str) -> str:
    if not any(_FULLWIDTH_START <= ord(char) <= _FULLWIDTH_END for char in text):
        return text.replace(_IDEOGRAPHIC_SPACE, " ")
    folded = []
    for char in text:
        code = ord(char)
        if _FULLWIDTH_START <= code <= _FULLWIDTH_END:
            folded.append(chr(code - _FULLWIDTH_OFFSET))
        elif char == _IDEOGRAPHIC_SPACE:
            folded.append(" ")
        else:
            folded.append(char)
    return "".join(folded)


def clean_text(text: str) -> str:
    """规整文本：换行、不可见字符、控制字符、全角 ASCII、行尾空白。

    只做**不改变语义**的规整。真正要判断"这段该不该留"的东西（页眉页脚）在
    ``extract_pdf`` 里做，因为那需要跨页信息，纯文本这一层看不到。
    """
    if not text:
        return ""
    normalized = text.replace("\r\n", "\n").replace("\r", "\n")
    normalized = normalized.translate(_INVISIBLE)
    normalized = _CONTROL_RE.sub("", normalized)
    normalized = _fold_fullwidth(normalized)
    normalized = _TRAILING_WS_RE.sub("", normalized)
    # 三个以上连续空行压成一个空行。_parse_blocks 遇到任何空行都会 flush 段落，
    # 所以这一步不改变分块结果，纯粹是别把空行当 token 存进库。
    normalized = _BLANK_RUN_RE.sub("\n\n", normalized)
    return normalized.strip()


def readable_ratio(text: str) -> float:
    """可读字符占比。入库自检用它判断"这份文档到底解出东西来了吗"。

    分母是非空白字符总数：一份正常文档这个数接近 1.0，一份被 ``errors="replace"``
    毁掉的 GBK 文档接近 0。返回 1.0 而不是 0.0 给空文本——空是另一种失败
    （``chunks == 0``），由调用方单独判，混在一起会让两种原因分不开。
    """
    meaningful = [char for char in text if not char.isspace()]
    if not meaningful:
        return 1.0
    bad = sum(1 for char in meaningful if char == _REPLACEMENT)
    return 1.0 - bad / len(meaningful)


# ========== 解码 ==========


@dataclass(frozen=True, slots=True)
class DecodeResult:
    text: str
    encoding: str
    replacement_ratio: float


def sniff_decode(content: bytes) -> DecodeResult:
    """嗅探编码并解码。顺序是 utf-8 严格解 → 配置的编码先验 → 嗅探器。

    **先严格试 utf-8**：utf-8 的多字节序列有自校验性，别的编码极难恰好通过，
    所以"能严格解通"基本等于"就是 utf-8"。反过来嗅探器对短文本会猜错，一份两行的
    utf-8 中文很容易被判成 GB18030，而后者也能解通——解出来是乱码但**不报错**。

    **然后才是配置的先验，最后才是嗅探器**，这一层是被实测逼出来的：一段 GBK
    编码的中文短句会被 charset-normalizer 判成 EUC-KR，解出来是一串谚文。这不是
    嗅探器的 bug——同一串字节在 GB18030 / EUC-KR / Shift-JIS 下**都能严格解通**，
    从字节本身无法分辨，任何检测器都不行。所以这里必须有一个先验，而先验属于配置。

    这个顺序的代价要说清楚：一份真正的韩文文档会被当成中文解错。对这个项目
    （中文语料、中文用户）这是正确的取舍，但它是取舍，不是正确性。

    最危险的一点：猜错编码时解出来的文本**没有任何替换符**，所以
    ``readable_ratio`` 抓不到它——那道自检只挡得住 ``errors="replace"`` 那一类。
    """
    if not content:
        return DecodeResult("", "utf-8", 0.0)

    try:
        # utf-8-sig 顺手吃掉 BOM。留着 BOM 的话它会变成正文第一个字符，
        # 而首行往往是 Markdown 标题，`﻿# 标题` 匹配不上 _HEADING_RE。
        return DecodeResult(content.decode("utf-8-sig"), "utf-8", 0.0)
    except UnicodeDecodeError:
        pass

    for encoding in _encoding_hints():
        try:
            return DecodeResult(content.decode(encoding), encoding, 0.0)
        except (UnicodeDecodeError, LookupError):
            continue

    try:
        from charset_normalizer import from_bytes

        best = from_bytes(content).best()
    except Exception as exc:  # 嗅探器自身出问题不该让上传挂掉
        logger.warning("charset sniffing failed: %s", type(exc).__name__)
        best = None

    if best is not None:
        text = str(best)
        return DecodeResult(text, str(best.encoding or "unknown"), readable_ratio(text))

    # 兜底还是 replace，但这次调用方拿得到 ratio，能据此把文档判成 failed
    # 而不是标成 indexed 然后永远检索不到。
    text = content.decode("utf-8", errors="replace")
    return DecodeResult(text, "utf-8/replace", readable_ratio(text))


def _encoding_hints() -> list[str]:
    """按配置给出严格解码的候选顺序。

    默认只有 ``gb18030``：它是 GBK / GB2312 的超集，能解通那两者的一切字节，
    所以一条就够，列三条只是让同一件事有三个出口。
    """
    return [
        item.strip()
        for item in (settings.INGEST_ENCODING_HINTS or "").split(",")
        if item.strip()
    ]


# ========== PDF 结构恢复 ==========


@dataclass(slots=True)
class PdfExtraction:
    text: str
    pages: int = 0
    backend: str = "pdfplumber"
    warnings: list[str] = field(default_factory=list)


@dataclass(slots=True)
class _Line:
    text: str
    size: float
    top: float
    right: float = 0.0  # 行末的 x 坐标，判断这行是不是被折行截断的
    full: bool = False  # 行末接近右边界 → 它是一段没写完的话
    level: int = 0  # 0 = 正文，1-6 = 标题层级


# 空格宽度大约是字号的四分之一。间隙比这个还小说明两段字符本来是连着的，
# 是按字距被切开的——那时补一个空格就等于凭空在词里插空格，而 BM25 的词元
# 正是按空格和字符类别切的。这个比例偏小是故意的：宁可漏补一个空格
# （`fooBar` 仍是一个词元），也不要在 `RESOURCE_EXHAUSTED` 中间插一个。
_SPACE_GAP_RATIO = 0.28
# 同一行的判定容差（点）。PDF 里同一行的基线会有亚像素抖动
_LINE_TOLERANCE = 2.5
# 字号超过正文这个倍数才算标题。1.15 是经验值：更小会把加粗的正文行误判成标题
_HEADING_SIZE_RATIO = 1.15
# 标题一般不长，也不以句末标点结尾。这两条是为了别把一整段大字号引言判成标题
_HEADING_MAX_CHARS = 80
_SENTENCE_END = "。！？；.!?;:："
# 行末 x 坐标达到本页最大行宽的这个比例，就认为这行是被折行截断的（见 _mark_full_lines）
_FULL_LINE_RATIO = 0.85
# 页眉页脚只在页面上下这个比例的区域里找
_MARGIN_ZONE = 0.10

_WORD_TAIL_RE = re.compile(r"[A-Za-z0-9_)\]}%.,;:!?'\"-]$")
_WORD_HEAD_RE = re.compile(r"^[A-Za-z0-9_(\[{'\"$#@]")
_DIGITS_RE = re.compile(r"\d+")
_LIST_HEAD_RE = re.compile(r"^\s*([-*+•]|\d+[.)、]|[（(]\d+[)）])\s")


def _needs_space(left: str, right: str) -> bool:
    """两段文本之间该不该加空格。

    只有拉丁/数字之间才需要。中文按字符切词时每个字都是一个 word，无条件用空格
    拼会把整篇中文变成「每 个 字 之 间 都 有 空 格」——而 ``tokenize()`` 的 CJK
    bigram 取的是**相邻字符**，插了空格 bigram 就全废了，稀疏通道对中文彻底失效。
    """
    if not left or not right:
        return False
    return bool(_WORD_TAIL_RE.search(left) and _WORD_HEAD_RE.match(right))


def _join_words(words: list[dict]) -> str:
    """把一行的词拼回文本，按几何间隙决定加不加空格。"""
    parts: list[str] = []
    previous: dict | None = None
    for word in words:
        text = str(word.get("text") or "")
        if not text:
            continue
        if previous is not None:
            gap = _as_float(word.get("x0")) - _as_float(previous.get("x1"))
            size = _as_float(word.get("size")) or _as_float(previous.get("size")) or 10.0
            if gap > size * _SPACE_GAP_RATIO and _needs_space(parts[-1], text):
                parts.append(" ")
        parts.append(text)
        previous = word
    return "".join(parts)


def _as_float(value: object) -> float:
    try:
        return float(value)  # type: ignore[arg-type]
    except (TypeError, ValueError):
        return 0.0


def _group_lines(words: list[dict]) -> list[_Line]:
    """按 ``top`` 把词分行。pdfplumber 给的是词，不是行。"""
    if not words:
        return []
    ordered = sorted(words, key=lambda word: (_as_float(word.get("top")), _as_float(word.get("x0"))))
    lines: list[_Line] = []
    bucket: list[dict] = []
    bucket_top = _as_float(ordered[0].get("top"))

    def flush() -> None:
        if not bucket:
            return
        text = _join_words(bucket).strip()
        if text:
            # 行字号取该行出现最多的那个，而不是最大的：一行里夹一个大号
            # 首字母或角标不该把整行判成标题
            sizes = Counter(round(_as_float(word.get("size")), 1) for word in bucket)
            lines.append(
                _Line(
                    text=text,
                    size=sizes.most_common(1)[0][0],
                    top=bucket_top,
                    right=max(_as_float(word.get("x1")) for word in bucket),
                )
            )
        bucket.clear()

    for word in ordered:
        top = _as_float(word.get("top"))
        if bucket and abs(top - bucket_top) > _LINE_TOLERANCE:
            flush()
            bucket_top = top
        bucket.append(word)
    flush()
    _mark_full_lines(lines)
    return lines


def _mark_full_lines(lines: list[_Line]) -> None:
    """标出"行末顶到右边界"的行。

    这是段落合并唯一可靠的信号。PDF 里没有段落标记，每一个折行都是一条独立的
    line，所以只靠"末尾没有句末标点"去合并会把**整页文字粘成一段**，顺带毁掉
    列表项和表格行。而排版规律是硬的：一段话被折行时前一行必然写到了右边界，
    一段话结束时最后一行通常留白。

    右边界取本页所有行 ``right`` 的最大值——PDF 里没有"页面正文宽度"这个元数据，
    只能从实际排版反推。
    """
    if not lines:
        return
    margin = max(line.right for line in lines)
    if margin <= 0:
        return
    for line in lines:
        line.full = line.right >= margin * _FULL_LINE_RATIO


# 数字之外还剩多少个"实字"才不算页码。页码是短模板 + 一个变化的数字：
# ``第 1 页`` 残余 2、``- 5 -`` 残余 0、``Page 12 of 30`` 残余 6。而位置在边缘
# 但仅数字不同的正文残余明显更长：``各页不同的边缘内容 1`` 是 9、
# ``表 1 显示了各部门的人员编制情况`` 是 14。取 6 两侧各留 2-3 字符余量。
_PAGE_NUMBER_RESIDUAL_MAX = 6

# 计算残余时要去掉的装饰字符。页码常写成 ``- 5 -``、``第 3 页 / 共 10 页``，
# 这些标记不承载语义，算进残余会把短页码误判成正文。
_DECORATION_RE = re.compile(r"[\s\-—–.,:：、（）()\[\]/|#]")


def _digit_residual(text: str) -> int:
    """去掉数字与装饰字符后还剩几个实字。"""
    return len(_DECORATION_RE.sub("", _DIGITS_RE.sub("", text)))


def _frequency_key(text: str) -> str:
    """页眉页脚的频率统计键。

    页码每页都不同，直接按原文统计频率永远只有 1 次、一条都剔不掉。折掉数字之后
    ``第 1 页`` 与 ``第 12 页`` collapse 成同一个键，才数得出"这东西每页都有"。

    但**只在这一行像页码时才折**。无条件折数字会把"位置在边缘、内容各页不同、
    差异恰好只在数字上"的正文也 collapse 成同一个键——``各页不同的边缘内容
    1/2/3`` 会变成一个键、计数 3、被当成页脚剔掉。这两个要求（剔页码 / 留正文）
    在纯数字折叠下是互斥的，判据只能来自别处。

    这里用"数字之外还剩几个实字"：页码是短模板 + 一个变化的数字，正文不是。
    残余超过阈值就退回精确匹配——那样它只有跨页**逐字重复**才会被剔，
    而逐字重复的边缘行本来就该当页眉处理（``公司内部资料`` 走的就是这条）。
    """
    stripped = text.strip()
    if _digit_residual(stripped) <= _PAGE_NUMBER_RESIDUAL_MAX:
        return _DIGITS_RE.sub("#", stripped)
    return stripped


def _strip_running_heads(pages: list[list[_Line]], heights: list[float]) -> int:
    """剔除页眉页脚。返回剔掉的行数。

    判据是**跨页重复**，不是位置——只按位置剔会把首页正文的第一行也剔掉。
    单页文档不做：一页看不出什么叫"重复出现"。
    """
    if len(pages) < max(2, settings.INGEST_HEADER_FOOTER_MIN_PAGES):
        return 0

    zone_counts: Counter[str] = Counter()
    for lines, height in zip(pages, heights):
        if height <= 0:
            continue
        margin = height * _MARGIN_ZONE
        seen: set[str] = set()
        for line in lines:
            if line.top <= margin or line.top >= height - margin:
                seen.add(_frequency_key(line.text))
        zone_counts.update(seen)

    threshold = max(2, settings.INGEST_HEADER_FOOTER_MIN_PAGES)
    repeated = {key for key, count in zone_counts.items() if count >= threshold}
    if not repeated:
        return 0

    removed = 0
    for index, (lines, height) in enumerate(zip(pages, heights)):
        if height <= 0:
            continue
        margin = height * _MARGIN_ZONE
        kept = []
        for line in lines:
            in_zone = line.top <= margin or line.top >= height - margin
            if in_zone and _frequency_key(line.text) in repeated:
                removed += 1
                continue
            kept.append(line)
        pages[index] = kept
    return removed


def _assign_heading_levels(pages: list[list[_Line]]) -> int:
    """按字号给行标标题层级。返回识别出的标题行数。

    正文字号取**按字符数加权**的众数，不是按行数：一份文档里标题行数可能比某个
    小字号脚注的行数还多，按行数投票会把正文认成标题。
    """
    weighted: Counter[float] = Counter()
    for lines in pages:
        for line in lines:
            weighted[line.size] += len(line.text)
    if not weighted:
        return 0

    body_size = weighted.most_common(1)[0][0]
    candidate_sizes = sorted(
        {
            line.size
            for lines in pages
            for line in lines
            if line.size > body_size * _HEADING_SIZE_RATIO
            and len(line.text) <= _HEADING_MAX_CHARS
            and line.text[-1] not in _SENTENCE_END
        },
        reverse=True,
    )
    if not candidate_sizes:
        return 0

    # 最大的字号是 #，往下依次 ##、###……超过 6 级都算 ######
    level_of = {size: min(6, index + 1) for index, size in enumerate(candidate_sizes)}
    count = 0
    for lines in pages:
        for line in lines:
            level = level_of.get(line.size)
            if (
                level is not None
                and len(line.text) <= _HEADING_MAX_CHARS
                and line.text[-1] not in _SENTENCE_END
            ):
                line.level = level
                count += 1
    return count


def _render(pages: list[list[_Line]]) -> str:
    """把分好级的行渲染成 Markdown，并把被折行/翻页切断的段落接回去。

    合并的条件是三个一起成立：上一行顶到了右边界（``full``，说明那句话没写完）、
    上一块不是标题、当前行不是列表项或新标题。少任何一条都会过度合并——只看
    "没有句末标点"会把整页粘成一段，连列表和表格行一起。
    """
    blocks: list[str] = []
    previous_full = False
    for lines in pages:
        for line in lines:
            if line.level:
                blocks.append(f"{'#' * line.level} {line.text}")
                previous_full = False
                continue
            if (
                blocks
                and previous_full
                and not blocks[-1].startswith("#")
                and blocks[-1][-1] not in _SENTENCE_END
                and not _LIST_HEAD_RE.match(line.text)
            ):
                glue = " " if _needs_space(blocks[-1], line.text) else ""
                blocks[-1] = f"{blocks[-1]}{glue}{line.text}"
            else:
                blocks.append(line.text)
            previous_full = line.full
    return "\n\n".join(blocks)


def structure_backend_available() -> bool:
    """pdfplumber 装上了吗。启动校验用它。

    单独暴露一个函数而不是让调用方 try import：缺依赖和"这个 PDF 解不了"是两回事，
    前者会让**每一份** PDF 都静默退回无结构抽取——同一份文档在两台机器上切出不同
    的块，而 eval 的结论就依赖这个。所以它必须在启动时就拦住，不能等上传时才发现。
    """
    try:
        import pdfplumber  # noqa: F401
    except ImportError:
        return False
    return True


def extract_pdf(content: bytes) -> PdfExtraction:
    """抽取 PDF 并恢复结构。失败时回退 PyPDF2 并记一条 warning。

    回退是运行时容错，不是依赖可选：pdfplumber 是必需依赖（``main.py`` 启动时
    校验），但单个畸形 PDF 让它抛异常时，退回纯文本抽取仍然比整个上传 500 好。
    """
    warnings: list[str] = []
    try:
        import pdfplumber
    except ImportError:
        # 走到这里说明启动校验被绕过了（比如直接调用脚本）。记一条显眼的
        # warning，让报告里能看出"这批文档根本没走结构恢复"。
        logger.error("pdfplumber not installed; PDF structure recovery is OFF")
        warnings.append("pdfplumber_missing")
        return extract_pdf_plain(content, warnings)

    try:
        page_lines: list[list[_Line]] = []
        heights: list[float] = []
        with pdfplumber.open(io.BytesIO(content)) as pdf:
            for page in pdf.pages:
                words = page.extract_words(extra_attrs=["size"]) or []
                page_lines.append(_group_lines(words))
                heights.append(_as_float(page.height))
    except Exception as exc:
        logger.warning("pdfplumber extraction failed: %s", type(exc).__name__)
        warnings.append(f"pdfplumber_failed:{type(exc).__name__}")
        return extract_pdf_plain(content, warnings)

    pages = len(page_lines)
    if not any(page_lines):
        # 有页面但一个词都没抽到：几乎一定是扫描件（只有图像层）。
        # 这不是异常，所以必须靠 warning + 入库自检兜住，否则它会静默变成
        # 一篇 chunks=0 的 indexed 文档。
        warnings.append("no_text_layer")
        return PdfExtraction(text="", pages=pages, warnings=warnings)

    removed = _strip_running_heads(page_lines, heights)
    if removed:
        warnings.append(f"running_heads_removed:{removed}")

    headings = _assign_heading_levels(page_lines)
    if headings:
        warnings.append(f"headings_recovered:{headings}")
    else:
        # 没有标题就意味着 chunking 的 heading_path 仍然恒为空，
        # 「章节边界优先」对这份文档依旧不生效。值得记下来。
        warnings.append("no_headings_detected")

    return PdfExtraction(
        text=clean_text(_render(page_lines)), pages=pages, warnings=warnings
    )


def extract_pdf_plain(content: bytes, warnings: list[str]) -> PdfExtraction:
    """PyPDF2 纯文本抽取。``INGEST_PDF_STRUCTURE=false`` 的对照组，也是回退路径。"""
    try:
        from PyPDF2 import PdfReader
    except ImportError:
        raise ValueError("PDF 解析需要安装 PyPDF2 库")

    try:
        reader = PdfReader(io.BytesIO(content))
        text = "\n\n".join(page.extract_text() or "" for page in reader.pages)
        pages = len(reader.pages)
    except Exception as exc:
        # 抛 ValueError 而不是让原始异常穿出去：路由把 ValueError 转成 400、
        # 其它异常转成 500，而"这个文件不是有效的 PDF"是用户错误不是服务错误。
        logger.warning("pypdf2 extraction failed: %s", type(exc).__name__)
        raise ValueError("无法解析该 PDF，文件可能已损坏或加了密") from exc

    if not text.strip():
        warnings.append("no_text_layer")
    return PdfExtraction(
        text=clean_text(text) if settings.INGEST_CLEAN else text,
        pages=pages,
        backend="pypdf2",
        warnings=warnings,
    )


# ========== .docx ==========
# 渲染成 Markdown，而不是新开一条结构路径。
#
# 理由是 chunking 那边已经有两套机制吃 Markdown 标题：给每块加标题路径
# （contextual chunk）、以及章节边界优先于填满 max_tokens。而
# ``_looks_like_markdown`` 在扩展名不匹配时会回退扫前 200 行找 `#`，所以一份
# 渲染出标题的 .docx 会自动走上那条路——不需要给 chunking 传任何新参数。
#
# 这也和 PDF 那条路径一致：``extract_pdf`` 同样是从字号恢复层级、发出 `#` 标题。
# docx 反而更简单，层级在段落样式名里是显式的，不必猜。

# Word 的标题样式名。中文版 Word 存的是本地化名字，两种都认——只认英文的话
# 一份中文 Word 文档会一个标题都识别不出来，然后静默退化成"无标题的长文本"，
# 症状就是 heading_path 恒为空、章节边界优先不生效。
_DOCX_HEADING_STYLES = ("heading", "标题")


def _docx_heading_level(style_name: str) -> int:
    """从段落样式名取标题层级。0 表示正文。

    形状是 ``Heading 1`` / ``标题 1``。``Title`` 与 ``Subtitle`` 不在内：
    它们是封面元素，映射成 h1 会让整篇文档挂在一个"标题"下面，等于没有层级。
    """
    lowered = (style_name or "").strip().lower()
    for prefix in _DOCX_HEADING_STYLES:
        if lowered.startswith(prefix):
            tail = lowered[len(prefix):].strip()
            if tail.isdigit():
                return min(6, max(1, int(tail)))
    return 0


def _docx_render_table(table: Any) -> str:
    """把表格渲染成 Markdown 表格。

    保留表格形状而不是拉平成句子：表头是**每一行都需要的上下文**，拉平之后
    "30 天" 这种单元格会脱离它的列名，检索命中了也答不出来。Markdown 表格能让
    表头和数据行留在同一个块里（分块按行不按字符切），也让模型看得懂列关系。

    xlsx 会走另一套（每行渲染成 `列名: 值`），那是块 4 的事：Excel 的一张表
    动辄几百行，Markdown 表格会被 max_tokens 拦腰截断，反而丢掉表头。
    """
    rows: list[list[str]] = []
    for row in table.rows:
        cells = [" ".join(cell.text.split()) for cell in row.cells]
        if any(cells):
            rows.append(cells)
    if not rows:
        return ""

    width = max(len(row) for row in rows)
    rows = [row + [""] * (width - len(row)) for row in rows]
    lines = ["| " + " | ".join(rows[0]) + " |"]
    lines.append("| " + " | ".join("---" for _ in range(width)) + " |")
    for row in rows[1:]:
        lines.append("| " + " | ".join(row) + " |")
    return "\n".join(lines)


def extract_docx(content: bytes) -> PdfExtraction:
    """抽取 .docx：段落 + 表格，标题渲染成 Markdown。

    复用 ``PdfExtraction`` 而不是新建一个 dataclass：字段需求完全一致
    （text / backend / warnings），而 ``pages`` 对 docx 没有意义所以留 0——
    Word 的分页是渲染期决定的，文件里没有这个信息。为此加一个几乎相同的
    类型，只会让 ``parse_document`` 里多一条分支去合并两种返回值。

    ``warnings`` 沿用 PDF 那套词汇（``headings_recovered:N`` /
    ``no_headings_detected``），这样入库自检和界面不必分格式各写一套判断。
    """
    warnings: list[str] = []
    try:
        import docx
    except ImportError:
        raise ValueError("Word 文档解析需要安装 python-docx 库")

    try:
        document = docx.Document(io.BytesIO(content))
    except Exception as exc:
        # 同 extract_pdf_plain：抛 ValueError 让路由转 400。
        # "这不是一个有效的 docx"是用户错误，不是服务错误。
        # 典型输入是把 .doc（老二进制格式）改名成 .docx——那读不了，
        # 而错误文案必须说清楚，否则用户会反复重传同一个文件。
        logger.warning("docx extraction failed: %s", type(exc).__name__)
        raise ValueError(
            "无法解析该 Word 文档。请确认它是 .docx（Word 2007 以后的格式）；"
            "老的 .doc 需要先另存为 .docx"
        ) from exc

    blocks: list[str] = []
    headings = 0
    tables = 0

    # 按文档流顺序遍历 body，而不是先 paragraphs 再 tables。
    # python-docx 的 ``document.paragraphs`` 与 ``document.tables`` 是两个独立
    # 列表，分别遍历会把所有表格搬到文末——表格于是脱离了它所属的小节，
    # 而标题路径正是按出现顺序算的。那样"报销标准"那张表会挂到最后一个标题下面。
    for child in document.element.body.iterchildren():
        tag = child.tag.split("}")[-1]
        if tag == "p":
            paragraph = docx.text.paragraph.Paragraph(child, document)
            text = paragraph.text.strip()
            if not text:
                continue
            level = _docx_heading_level(
                paragraph.style.name if paragraph.style is not None else ""
            )
            if level:
                headings += 1
                blocks.append("#" * level + " " + text)
            else:
                blocks.append(text)
        elif tag == "tbl":
            table = docx.table.Table(child, document)
            rendered = _docx_render_table(table)
            if rendered:
                tables += 1
                blocks.append(rendered)

    if headings:
        warnings.append(f"headings_recovered:{headings}")
    else:
        # 同 PDF：没有标题就意味着 heading_path 恒为空、章节边界优先不生效。
        # 对 Word 来说这通常是"作者用加粗大字号代替了标题样式"——很常见，
        # 而它对检索质量的影响是实打实的，所以必须留痕。
        warnings.append("no_headings_detected")
    if tables:
        warnings.append(f"tables_extracted:{tables}")
    if not blocks:
        # 有文件、但一个段落一张表都没抽到。对应 PDF 的 no_text_layer：
        # 几乎一定是整篇内容都在文本框/图片里，而那两种 python-docx 读不到。
        # 不抛异常——它是一篇合法的 docx，只是对检索没有价值，
        # 靠 warning + 入库自检兜住，别静默变成 chunks=0 的 indexed 文档。
        warnings.append("no_extractable_text")

    text = "\n\n".join(blocks)
    return PdfExtraction(
        text=clean_text(text) if settings.INGEST_CLEAN else text,
        pages=0,
        backend="python-docx",
        warnings=warnings,
    )


# ========== .xlsx ==========
# docx 的表格渲染成 Markdown 表格，xlsx **不能**照抄。这是块 4 唯一的真实取舍。
#
# 一张 Excel 表动辄几百行。渲染成 Markdown 表格的话：
#   - 表头只在第一行出现一次，而 max_tokens 默认 320，几十行就要断开；
#   - 断开之后从第二块起全是裸数据行，"500" 脱离了"住宿标准"这个列名。
#     检索命中了也答不出问题——这正是保留表格形状本来想避免的事。
#
# 所以每行**自带列名**渲染成 `列名: 值 | 列名: 值`：
#   - 每一行都是自足的，切在哪都不会丢上下文；
#   - 行之间用空行分隔 → ``_parse_blocks`` 遇空行 flush，每行成为独立 block；
#   - 行内不用句末标点（`：` 是全角冒号不在 ``_SENTENCE_RE`` 里，分隔用 `|`），
#     所以 ``_split_units`` 的句子切分不会从行中间切开。
#
# 代价是重复：列名在每一行都出现一次，token 数比 Markdown 表格高。
# 换来的是"任何一块都能独立回答问题"，对检索这是划算的——冗余的是提示词成本，
# 而丢上下文是答不出来。

# 一张表最多读多少行。不是性能考虑（read_only 是流式的），是**信噪比**：
# 上万行的明细表进知识库，检索时会用几百个近乎相同的块淹掉其它文档。
# 超出就截断并留 warning——静默截断是这个仓库里最不能接受的一类行为。
_XLSX_MAX_ROWS = 500
# 一张表最多读多少列。超宽表通常是把多张表拼在一起，取前 N 列已经够用。
_XLSX_MAX_COLS = 50


def _xlsx_cell_text(value: Any) -> str:
    """把单元格值转成文本。

    ``datetime`` 单独处理：默认的 ``str()`` 会给出 ``2026-08-24 00:00:00``，
    那个恒为零的时间部分是噪声，也会让 BM25 多出一堆无意义词元。
    ``float`` 的整数值同理——``100.0`` 应当是 ``100``，否则"每日 100 元"这类
    查询字面匹配不上。
    """
    if value is None:
        return ""
    if isinstance(value, datetime):
        if value.hour or value.minute or value.second:
            return value.strftime("%Y-%m-%d %H:%M:%S")
        return value.strftime("%Y-%m-%d")
    if isinstance(value, date):
        return value.strftime("%Y-%m-%d")
    if isinstance(value, float) and value.is_integer():
        return str(int(value))
    return " ".join(str(value).split())


def _xlsx_header_row(rows: list[list[str]]) -> tuple[list[str], int]:
    """挑出表头行，返回 ``(列名, 数据起始下标)``。

    不是无脑取第一行：真实的表格常常第一行是标题（"2026 年差旅标准"，只占 A1，
    后面全空），表头在第二行。判据是"非空单元格最多的那一行"，只在前 5 行里找。

    找错的后果不是报错，是每一行都带上错的列名——一整篇文档静默变成噪声。
    所以宁可用一个能解释的启发式，也不要"第一行就是表头"这个会安静出错的假设。
    """
    limit = min(5, len(rows))
    best_index = 0
    best_filled = -1
    for index in range(limit):
        filled = sum(1 for cell in rows[index] if cell)
        # 严格大于：并列时取更靠上的那一行，表头一般在数据之上
        if filled > best_filled:
            best_filled = filled
            best_index = index
    header = [cell or f"第{position + 1}列" for position, cell in enumerate(rows[best_index])]
    return header, best_index + 1


def extract_xlsx(content: bytes) -> PdfExtraction:
    """抽取 .xlsx：每张工作表一个二级标题，每行一个自足的 ``列名: 值`` 块。

    复用 ``PdfExtraction``（理由同 ``extract_docx``），``pages`` 借用来记工作表数——
    这是它对 Excel 唯一说得通的含义。
    """
    warnings: list[str] = []
    try:
        import openpyxl
    except ImportError:
        raise ValueError("Excel 文档解析需要安装 openpyxl 库")

    try:
        # data_only=True 取公式的**缓存值**而不是 "=SUM(B2:B9)" 这个字符串。
        #
        # 已知限制（没有修，代价不划算）：文件如果从未被 Excel 打开过（程序生成的
        # 就是这样），公式单元格没有缓存值，取到 None——那一列会整列变空，
        # 而这里区分不出"空单元格"和"公式没缓存值"：data_only 模式下 openpyxl
        # 直接丢掉了公式本身。要分辨就得把工作簿加载两遍（data_only=False 看公式、
        # True 看值），对几万行的表是双倍开销。
        # 症状：那一列检索不到，而文档 indexed、chunks 非零。
        # 真遇到时的处置是让用户用 Excel 打开另存一次。
        workbook = openpyxl.load_workbook(
            io.BytesIO(content), read_only=True, data_only=True
        )
    except Exception as exc:
        logger.warning("xlsx extraction failed: %s", type(exc).__name__)
        raise ValueError(
            "无法解析该 Excel 文档。请确认它是 .xlsx（Excel 2007 以后的格式）；"
            "老的 .xls 需要先另存为 .xlsx"
        ) from exc

    blocks: list[str] = []
    sheets = 0
    total_rows = 0
    truncated: list[str] = []

    try:
        for sheet in workbook.worksheets:
            raw: list[list[str]] = []
            for position, row in enumerate(sheet.iter_rows(values_only=True)):
                if position >= _XLSX_MAX_ROWS:
                    truncated.append(sheet.title)
                    break
                cells = [_xlsx_cell_text(value) for value in row[:_XLSX_MAX_COLS]]
                if any(cells):
                    raw.append(cells)
            if not raw:
                continue

            sheets += 1
            # 工作表名当二级标题：它是这批行的共同上下文（"住宿标准" vs "交通标准"），
            # 而 chunking 会把它放进每一块的 heading_path。
            blocks.append(f"## {sheet.title}")

            header, start = _xlsx_header_row(raw)
            width = len(header)
            for cells in raw[start:]:
                padded = cells + [""] * (width - len(cells))
                # 只保留非空单元格：空值渲染成 `列名: ` 是纯噪声，
                # 而稀疏表（很多可选列）在真实数据里很常见。
                pairs = [
                    f"{header[position]}: {value}"
                    for position, value in enumerate(padded[:width])
                    if value
                ]
                if pairs:
                    total_rows += 1
                    blocks.append(" | ".join(pairs))
    finally:
        # read_only 模式持有文件句柄，不关会在 Windows 上把临时文件锁住
        workbook.close()

    if sheets:
        warnings.append(f"sheets_extracted:{sheets}")
        warnings.append(f"rows_extracted:{total_rows}")
    if truncated:
        # 截断必须可见。静默丢掉后半张表的症状是"某些条目怎么都检索不到"，
        # 而文档状态是 indexed、chunks 也非零，看起来一切正常。
        warnings.append(
            f"rows_truncated:{_XLSX_MAX_ROWS}:{','.join(sorted(set(truncated)))}"
        )
    if not blocks:
        warnings.append("no_extractable_text")

    text = "\n\n".join(blocks)
    return PdfExtraction(
        text=clean_text(text) if settings.INGEST_CLEAN else text,
        pages=sheets,
        backend="openpyxl",
        warnings=warnings,
    )
